#!/usr/bin/env python3
"""Generate the internal strategy doc for AI User-Facing Docs project."""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

doc = Document()

style = doc.styles['Normal']
style.font.name = 'Arial'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

for level in range(1, 4):
    hstyle = doc.styles[f'Heading {level}']
    hstyle.font.name = 'Arial'
    hstyle.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
    return table

# ============================================================
# TITLE
# ============================================================
title = doc.add_heading('AI-Generated User-Facing Documentation', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.LEFT

subtitle = doc.add_paragraph()
run = subtitle.add_run('Automating OpenShift/OKD documentation using LLM skills trained on source code diffs')
run.font.size = Pt(13)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

meta = doc.add_paragraph()
meta.add_run('Team: ').bold = True
meta.add_run('OKD Engineering  |  ')
meta.add_run('Author: ').bold = True
meta.add_run('Sarthak Purohit  |  ')
meta.add_run('Date: ').bold = True
meta.add_run('September 2026')

meta2 = doc.add_paragraph()
meta2.add_run('Repository: ').bold = True
meta2.add_run('https://github.com/sarthakpurohit/ai-user-docs')
meta2.add_run('  |  ')
meta2.add_run('Presentation: ').bold = True
meta2.add_run('presentation/project-overview-final.html (in repo)')

doc.add_paragraph()

# ============================================================
# TL;DR
# ============================================================
doc.add_heading('tl;dr', level=1)

p = doc.add_paragraph()
p.add_run('WHAT: ').bold = True
p.add_run(
    'An AI pipeline that generates updated OpenShift/OKD documentation for each release. '
    'It reads the previous version\'s docs, analyzes code changes across multiple source repositories, '
    'and produces the new version\'s docs. A trained "SKILL.md" file encodes documentation conventions '
    'and lessons learned from prior iterations.'
)

p = doc.add_paragraph()
p.add_run('RESULTS: ').bold = True
p.add_run(
    '94 to 97% text similarity to human-written docs on stable releases. '
    'Trained on 2 documentation sections (Installing, Updating) across 6 version iterations (4.17 through 4.22).'
)

p = doc.add_paragraph()
p.add_run('IMPACT: ').bold = True
p.add_run(
    'The docs team gets a high-quality first draft within one day of feature freeze, '
    'months before GA. Instead of days of manual writing per section per release, '
    'the effort shifts to focused review and editorial polish.'
)

p = doc.add_paragraph()
p.add_run('NEXT: ').bold = True
p.add_run(
    'Ready for production use on 4.23 and beyond. '
    'Can be extended to more documentation sections (networking, security, observability) using the same approach. '
    'Adding JIRA/epic context as an additional input could close another 10 to 15% of the accuracy gap.'
)

doc.add_paragraph()

# ============================================================
# 1. THE PROBLEM
# ============================================================
doc.add_heading('1. The Problem', level=1)

doc.add_paragraph(
    'OpenShift documentation is built from the openshift/openshift-docs repository using AsciiDoc. '
    'It covers about 20 sections (Installing, Updating, Networking, Security, etc.) that must be '
    'refreshed every release. OKD, the upstream community distribution, shares the same documentation '
    'source, differentiated via AsciiDoc conditionals.'
)

bullets = [
    'The central docs team is withdrawing dedicated support for OKD documentation.',
    'OpenShift targets a 4-month release cadence (3 releases per year). Manual doc writing cannot keep up.',
    'Existing OKD docs are outdated, still referencing Fedora CoreOS and incorrect operators.',
    'Each release touches 50 to 200 user-facing changes across multiple source repositories.',
]
for b in bullets:
    doc.add_paragraph(b, style='List Bullet')

p = doc.add_paragraph()
run = p.add_run('The gap: ')
run.bold = True
run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
p.add_run(
    'Without automation, documenting each release requires significant manual effort from an '
    'already stretched docs team. Code changes happen across many repositories, and keeping docs '
    'in sync takes days of work per release cycle.'
)

# ============================================================
# 2. THE APPROACH
# ============================================================
doc.add_heading('2. The Approach', level=1)

doc.add_paragraph('One formula drives the entire system:')

formula = doc.add_paragraph()
formula.alignment = WD_ALIGN_PARAGRAPH.CENTER
formula.paragraph_format.space_before = Pt(12)
formula.paragraph_format.space_after = Pt(12)
run = formula.add_run('Docs(4.x)  =  SKILL.md  +  Docs(4.x-1)  +  CodeDiff(4.x-1 to 4.x)')
run.bold = True
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x00, 0x66, 0xCC)

doc.add_paragraph(
    'In plain English: give the LLM a trained skill file, last version\'s docs, and what changed in '
    'the source code. It produces the new version\'s docs.'
)

doc.add_heading('How it works step by step', level=2)

steps = [
    ('Extract', 'Pull baseline docs for a known version from openshift/openshift-docs.'),
    ('Diff', 'Generate structured code diffs between two release branches across the relevant source repos.'),
    ('Generate', 'Feed SKILL.md, baseline docs, and code diff to an LLM.'),
    ('Score', 'Compare generated output against ground truth using deterministic metrics.'),
    ('Evaluate', 'Run LLM-based semantic evaluation for deeper quality assessment.'),
    ('Learn', 'Identify what the LLM got wrong, add corrective rules to SKILL.md.'),
    ('Repeat', 'Move to the next version. Each pass adds 2 to 5 battle-tested rules.'),
]
for name, desc in steps:
    p = doc.add_paragraph()
    p.add_run(f'{name}: ').bold = True
    p.add_run(desc)

doc.add_heading('What is a "Code Diff"?', level=2)

doc.add_paragraph(
    'Not a raw git diff. It is a structured, filtered summary of documentation-relevant code changes. '
    'Each section monitors a specific set of source repositories and extracts:'
)
diff_items = [
    'Full file contents of key Go type definition files (struct definitions, field comments, defaults).',
    'CRD schema sections with field hierarchy context.',
    'Filtered key changes (only lines matching doc-relevant patterns like type definitions, json tags, validation annotations).',
    'Commit messages that explain the intent behind changes.',
    'New and deleted file lists showing structural additions and removals.',
]
for d in diff_items:
    doc.add_paragraph(d, style='List Bullet')

doc.add_paragraph(
    'We discovered that "enhanced diffs," which include full Go type file contents and CRD schemas, '
    'improve LLM completeness by +34% compared to basic diffs showing only changed lines.'
)

# ============================================================
# 3. SOURCE REPOSITORIES
# ============================================================
doc.add_heading('3. Source Repositories', level=1)

doc.add_paragraph(
    'Each documentation section monitors a specific set of OpenShift source repositories for code changes. '
    'Different sections need different repos because the code that affects installation docs is different '
    'from the code that affects update docs. The repos below are what we identified and validated during '
    'training for the two sections covered so far.'
)

doc.add_heading('Installing section (7 repos)', level=2)

add_table(
    ['Repository', 'What it contributes to docs'],
    [
        ['openshift/installer', 'Install-config Go structs, platform types, CLI commands, CRD schema, UPI templates'],
        ['openshift/api', 'ClusterVersion types, feature gates, machine API types, network types'],
        ['openshift/baremetal-operator', 'Bare metal provisioning APIs, BMH (BareMetalHost) types'],
        ['openshift/assisted-installer', 'Assisted installer workflows, agent-based installation'],
        ['openshift/cluster-network-operator', 'Network plugin configuration, SDN/OVN settings'],
        ['openshift/machine-config-operator', 'Node configuration, MachineConfig types, OS-level settings'],
        ['openshift/machine-api-operator', 'Machine API types, provider-specific machine specs'],
    ]
)

doc.add_paragraph()
doc.add_heading('Updating section (5 repos)', level=2)

add_table(
    ['Repository', 'What it contributes to docs'],
    [
        ['openshift/cluster-version-operator', 'Update orchestration, preconditions, upgrade gates, status history'],
        ['openshift/oc', 'CLI commands: oc adm upgrade (recommend, status, channel, rollback)'],
        ['openshift/machine-config-operator', 'Node update behavior: drain logic, reboot, MCP coordination'],
        ['openshift/api', 'ClusterVersion API types, condition types, Kubernetes version (go.mod)'],
        ['openshift/cluster-network-operator', 'SDN/OVN migration gates, CNI changes that block updates'],
    ]
)

doc.add_paragraph()
doc.add_paragraph(
    'Note: machine-config-operator, api, and cluster-network-operator are shared between both sections. '
    'As more sections are trained, additional repos will be identified through the same iterative process.'
)

# ============================================================
# 4. THE SKILL FILE
# ============================================================
doc.add_heading('4. The Heart: SKILL.md', level=1)

doc.add_paragraph(
    'A SKILL.md file is a detailed instruction manual for the LLM. '
    'It tells the AI exactly what to look for in code diffs, what rules to follow when generating '
    'documentation, and what mistakes to avoid. Each documentation section has its own skill file, '
    'trained iteratively.'
)

add_table(
    ['Section', 'Skill File', 'Rules', 'Source Repos', 'Avg Text Similarity'],
    [
        ['Installing', 'skills/generate-install-docs/SKILL.md', '23 rules', '7 repos', '94.8%'],
        ['Updating', 'skills/generate-updating-docs/SKILL.md', '28 rules', '5 repos', '96.5%'],
    ]
)

doc.add_paragraph()
doc.add_paragraph('Example rules learned through iterative training:')
rules = [
    'Do not fabricate content. If a field is not in the diff or source code, do not document it.',
    'Go struct field comments map directly to parameter descriptions in docs.',
    'json:"fieldName" tags become the parameter names shown to users.',
    'When a feature gate is removed (GA promotion), the command becomes the primary approach.',
    'Do not bump admin-ack gate strings without explicit evidence in the diff.',
    'Single-agent generation is mandatory. Never split work into subagents.',
]
for r in rules:
    doc.add_paragraph(r, style='List Bullet')

# ============================================================
# 5. TRAINING LOOP
# ============================================================
doc.add_heading('5. The Iterative Training Loop', level=1)

doc.add_paragraph(
    'We improve the skill by testing it against versions where we already have human-written docs '
    'as ground truth. Versions 4.16 through 4.22 serve as training data.'
)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(8)
p.paragraph_format.space_after = Pt(8)
run = p.add_run('Generate 4.17  >  Score vs actual  >  Learn mistakes  >  Update SKILL.md  >  Generate 4.18  >  ...')
run.font.size = Pt(11)
run.bold = True
run.font.color.rgb = RGBColor(0x00, 0x66, 0xCC)

doc.add_paragraph(
    'After 6 iterations, the skill grows from about 12 generic rules to 23 to 28 specific rules '
    'covering edge cases discovered through real failures. Each pass adds 2 to 5 new rules.'
)

doc.add_heading('Skill evolution: Installing (12 to 23 rules)', level=3)
doc.add_paragraph(
    'Start: 12 rules. After 4.17: +3. After 4.18: +3. After 4.19: +1. '
    'After 4.20: +1. After 4.21: +2. After 4.22: +1. Final: 23 rules.'
)
doc.add_paragraph(
    'Key additions: struct-to-docs mapping, CRD schema handling, platform constraint rules, '
    'directory restructuring detection, Go comments as doc source.'
)

doc.add_heading('Skill evolution: Updating (12 to 28 rules)', level=3)
doc.add_paragraph(
    'Start: 12 rules. After 4.17: +5. After 4.18: +2. After 4.19: +3. '
    'After 4.20: +3. After 4.21: +1. After 4.22: +2. Final: 28 rules.'
)
doc.add_paragraph(
    'Key additions: admin-ack evidence logic, no-subagent rule, GA promotion handling, '
    'CLI fixture usage, file removal detection, Kubernetes version mapping.'
)

# ============================================================
# 6. EVALUATION METRICS
# ============================================================
doc.add_heading('6. How We Measure Quality', level=1)

doc.add_heading('Deterministic scoring (fast, automated)', level=2)
doc.add_paragraph('Run with "make score VERSION=4.17 SECTION=updating." Takes seconds.')

add_table(
    ['Metric', 'What it measures', 'Method'],
    [
        ['File Coverage', 'Are all expected files present?', 'Count files in generated vs ground truth'],
        ['Text Similarity', 'How close is the content line by line?', 'difflib.SequenceMatcher ratio (0 to 100%)'],
        ['Section Coverage', 'Are all section headings present?', 'Extract == Heading patterns, check presence'],
        ['Param Coverage', 'Are all parameter names mentioned?', 'Extract backtick-quoted terms, check presence'],
    ]
)

doc.add_paragraph()
doc.add_heading('LLM-based evaluation (deeper, semantic)', level=2)
doc.add_paragraph('Run in a separate agent window. Takes 10 to 20 minutes per version.')

add_table(
    ['Metric', 'What it measures', 'Method'],
    [
        ['Semantic Accuracy', 'Are the facts correct?', 'LLM reads each file, classifies as correct/minor/major issue'],
        ['Completeness', 'Are all code changes reflected?', 'LLM reads diff, checks each change is documented'],
        ['Structure Compliance', 'Are AsciiDoc/OpenShift conventions followed?', 'Check headers, IDs, includes, attributes'],
        ['Command Accuracy', 'Are CLI commands correct?', 'Verify against actual source code'],
    ]
)

doc.add_paragraph()
doc.add_paragraph(
    'Why two levels? Deterministic scoring catches surface problems like missing files or wrong version numbers. '
    'LLM evaluation catches semantic issues: whether the meaning is right, whether the documentation '
    'accurately describes what the code actually does.'
)

# ============================================================
# 7. RESULTS
# ============================================================
doc.add_heading('7. Results', level=1)

doc.add_heading('Installing section', level=2)
doc.add_paragraph('800 to 1000 files per version  |  7 source repos  |  6 iterations  |  23 rules')

add_table(
    ['Version', 'File Coverage', 'Text Similarity', 'Section Coverage', 'Param Coverage'],
    [
        ['4.17', '94.8%', '92.1%', '93.3%', '93.9%'],
        ['4.18', '95.7%', '94.2%', '94.5%', '95.1%'],
        ['4.19', '97.0%', '96.4%', '96.9%', '96.9%'],
        ['4.20 (major restructure)', '90.8%', '83.2%', '87.5%', '87.8%'],
        ['4.21', '97.3%', '96.4%', '97.2%', '97.0%'],
        ['4.22', '96.2%', '94.8%', '95.6%', '95.3%'],
    ]
)

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Average text similarity (excluding 4.20 restructure): ').bold = True
p.add_run('94.8%')

doc.add_paragraph()
doc.add_paragraph(
    'Version 4.20 is an outlier because the docs team restructured the Installing section that release, '
    'reorganizing files into new subdirectories. Since this was an editorial decision not visible in code, '
    'the AI could not anticipate it.'
)

doc.add_heading('LLM evaluation: Installing 4.17 (two rounds)', level=3)
doc.add_paragraph(
    'We discovered that giving the LLM richer code context dramatically improves quality:'
)

add_table(
    ['Metric', 'Before (basic diff)', 'After (enhanced diff)', 'Improvement'],
    [
        ['Semantic Accuracy', '91.2%', '92.7%', '+1.5%'],
        ['Completeness', '46.4%', '80.8%', '+34.4%'],
        ['Parameter Accuracy', '82.3%', '94.7%', '+12.4%'],
    ]
)

doc.add_paragraph()
doc.add_paragraph(
    'What changed: a "basic diff" shows only changed lines, like a typical git diff. '
    'An "enhanced diff" includes the full contents of key files: Go struct definitions with field comments, '
    'CRD schemas with field hierarchies, and CLI source code. '
    'This gives the LLM enough context to correctly name parameters, describe defaults, '
    'and write accurate procedure steps.'
)

doc.add_paragraph()
doc.add_heading('Updating section', level=2)
doc.add_paragraph('104 to 115 files per version  |  5 source repos  |  6 iterations  |  28 rules')

add_table(
    ['Version', 'File Coverage', 'Text Similarity', 'Section Coverage', 'Param Coverage'],
    [
        ['4.17', '100%', '96.4%', '96.5%', '99.1%'],
        ['4.18', '97.3%', '95.7%', '96.9%', '96.5%'],
        ['4.19', '100%', '96.6%', '100%', '98.7%'],
        ['4.20 (GA promotion + restructure)', '87.0%', '79.3%', '85.5%', '85.2%'],
        ['4.21', '99.1%', '98.8%', '99.1%', '99.1%'],
        ['4.22', '95.6%', '94.9%', '95.3%', '95.6%'],
    ]
)

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Average text similarity (excluding 4.20 restructure): ').bold = True
p.add_run('96.5%')

doc.add_paragraph()
doc.add_heading('LLM evaluation: Updating (across iterations)', level=3)

add_table(
    ['Version', 'LLM Overall', 'Key Issues Found'],
    [
        ['4.17 (run 1, 3 repos)', '80.8%', 'Fabricated admin-ack gate, wrong kubelet version'],
        ['4.17 (run 2, 5 repos)', '58%', 'Orphan files kept, SDN removal missed'],
        ['4.18', '65%', 'New recommend subcommand missed, ccoctl naming'],
        ['4.19', '48%', 'KMM restructure from non-monitored repo'],
        ['4.20', '52%', '15 new files (overview restructure), bootloader docs'],
        ['4.21', '62%', '1 missing file (new module), minor version mismatches'],
    ]
)

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Why are LLM scores lower than deterministic? ').bold = True
p.add_run(
    'Deterministic scoring measures surface similarity, and most files are copied unchanged, scoring 100%. '
    'LLM evaluation focuses on the 10 to 15 files that should have changed. '
    'Those files carry the hardest changes: product decisions, non-monitored repo changes, '
    'and editorial restructuring. The LLM judge penalizes every missed change, '
    'even if it is invisible in code diffs.'
)

doc.add_paragraph()
doc.add_heading('Cross-section comparison', level=2)

add_table(
    ['Metric', 'Installing', 'Updating'],
    [
        ['Files per version', '800 to 1000', '104 to 115'],
        ['Source repos monitored', '7', '5'],
        ['Avg text similarity (excl 4.20)', '94.8%', '96.5%'],
        ['Best text similarity', '96.4% (4.19, 4.21)', '98.8% (4.21)'],
        ['Worst version (4.20)', '83.2%', '79.3%'],
        ['Nature of content', 'Code-heavy (Go structs, CRDs)', 'Procedural (CLI, workflows)'],
    ]
)

doc.add_paragraph()
doc.add_paragraph(
    'The approach works across both section types. Deterministic scores are comparable, '
    'around 95 to 97% on stable releases. The LLM evaluation gap between sections reflects '
    'the nature of the content: installing docs are heavily code-derived (Go structs, CRDs), '
    'so the LLM covers most changes. Updating docs depend more on product decisions that are invisible in code.'
)

# ============================================================
# 8. KEY FINDINGS
# ============================================================
doc.add_heading('8. Key Findings', level=1)

doc.add_heading('Architecture findings', level=2)
findings = [
    ('Single agent is mandatory',
     'Subagents destroy consistency. When work is split into parallel tasks, each subagent loses context '
     'of the full rule set and other changes. The result is contradictory edits, fabricated content, '
     'and rule violations. Subagents scored 45% on structure versus 95 to 100% for a single agent.'),
    ('Enhanced diffs are essential',
     'Full Go type files give +34% completeness for installing. Full CVO/oc source gives accurate '
     'CLI docs for updating. The LLM needs to see the complete field definitions, not just the changed lines.'),
    ('Multiple repos are necessary',
     'Installing needs 7 repos, updating needs 5. No single repo has all the changes. '
     'We expanded from 3 to 7 repos during training for installing as gaps were discovered.'),
    ('Iterative training works',
     'Installing went from 12 to 23 rules. Updating went from 12 to 28 rules. '
     'Each pass fixes real failures and adds specific, battle-tested rules.'),
]
for title, desc in findings:
    p = doc.add_paragraph()
    p.add_run(f'{title}. ').bold = True
    p.add_run(desc)

doc.add_heading('Content findings', level=2)
content_findings = [
    ('Code comments are documentation',
     'Go struct field comments map directly to parameter descriptions. CLI help text maps to procedure steps.'),
    ('GA promotions change workflows',
     'When a feature gate is removed, the command becomes the primary approach (updating). '
     'When a tech preview platform becomes GA, it gets full procedure docs (installing).'),
    ('Version sweeps need evidence',
     'Do not bump version numbers, admin-ack strings, or API versions without explicit evidence in the diff.'),
    ('The agent is sometimes ahead of human writers',
     'In several cases, the AI documented real code changes that human-written docs had not yet caught up to. '
     'All verified against source code.'),
]
for title, desc in content_findings:
    p = doc.add_paragraph()
    p.add_run(f'{title}. ').bold = True
    p.add_run(desc)

# ============================================================
# 9. LIMITATIONS
# ============================================================
doc.add_heading('9. Limitations and the Accuracy Gap', level=1)

doc.add_paragraph(
    'About 20 to 25% of documentation changes between versions come from product or editorial decisions '
    'that live in no source repository. This is the hard ceiling for a code-diff-only approach.'
)

doc.add_heading('Accuracy gap breakdown', level=2)

add_table(
    ['Category', '% of Changes', 'Example', 'Why invisible in code'],
    [
        ['Code-derivable (handled)', '~75%', 'New struct fields, CLI flags, CRD changes', 'This IS what we handle'],
        ['Product decisions', '~15%', 'RHV deprecation, RHEL worker removal', 'Product management decision'],
        ['Non-monitored repos', '~10%', 'KMM restructure, ccoctl, Gateway API', 'Not in our current source repos'],
        ['Agent errors', '~5%', 'Fabricated content, stale ack strings', 'Fixable with better rules'],
    ]
)

doc.add_paragraph()
doc.add_heading('What code diffs cannot catch', level=2)

items = [
    'Content relocation: docs team decisions to reorganize files and directories.',
    'Product deprecation: product management decisions to remove platform support.',
    'Non-monitored repos: changes from repositories not in our current watch list.',
    'Editorial restructuring: writer preferences for assembly splits, title rewording.',
    'New content creation: entirely new documentation pages that cannot be derived from any code.',
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('Potential improvement: adding JIRA/Epic context', level=2)

doc.add_paragraph(
    'The core approach (SKILL.md + iterative training) stays the same. '
    'What changes is giving the LLM additional context beyond code diffs. '
    'Every release has tracked epics in JIRA that capture product decisions:'
)

add_table(
    ['Gap Category', 'How JIRA Helps'],
    [
        ['Product deprecations', 'Epics explicitly track deprecation decisions'],
        ['Non-monitored repos', 'Epics link to PRs across all repos, not just the ones we watch'],
        ['New content creation', 'Enhancement epics describe new features with acceptance criteria'],
        ['Content relocation', 'Docs-team epics may describe IA restructuring plans'],
    ]
)

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run('Estimated impact: ').bold = True
p.add_run(
    'Adding JIRA context could close 10 to 15% of the current gap, '
    'bringing automated coverage from about 75 to 80% up to roughly 90%, '
    'without changing the core formula.'
)

# ============================================================
# 10. PRODUCTION READINESS
# ============================================================
doc.add_heading('10. Production Readiness (4.23 and Beyond)', level=1)

doc.add_paragraph(
    'For future versions where there are no human-written docs to compare against, '
    'the pipeline generates a high-quality first draft.'
)

doc.add_heading('What the AI handles (about 75 to 80% of changes)', level=2)
ai_handles = [
    'New install-config fields from Go structs and CRD schemas.',
    'Platform constraint changes (validation rules).',
    'New CLI commands and flags (oc adm upgrade subcommands).',
    'Changed CLI output formats (from test fixtures).',
    'New CVO preconditions and upgrade gates.',
    'MCO behavioral changes (drain, reboot).',
    'Version string sweeps across all files.',
    'Correct AsciiDoc structure, includes, and conditionals.',
]
for a in ai_handles:
    doc.add_paragraph(a, style='List Bullet')

doc.add_heading('What humans still do (about 20 to 25% of changes)', level=2)
humans_do = [
    'Content reorganization decisions.',
    'New platform or feature documentation from scratch.',
    'Product deprecation and removal notices.',
    'Cross-section coordination (xrefs, includes).',
    'Editorial quality review.',
    'Final approval before publishing.',
]
for h in humans_do:
    doc.add_paragraph(h, style='List Bullet')

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('Bottom line: ')
run.bold = True
run.font.color.rgb = RGBColor(0x00, 0x88, 0x00)
p.add_run(
    'The pipeline generates a high-quality first draft that is 93 to 99% accurate on code-derivable changes. '
    'The docs team reviews, adds product context, and approves. '
    'Time savings: days of manual writing per section per release shift to focused review and editorial polish.'
)

# ============================================================
# 11. FUTURE WORKFLOW
# ============================================================
doc.add_heading('11. Future Workflow', level=1)

doc.add_paragraph(
    'Each trained SKILL.md becomes a self-contained doc generator for its section. '
    'The production workflow is straightforward:'
)

workflow_steps = [
    'Feature freeze happens (e.g., Aug 13 for 4.23).',
    'Run code diff generation across all monitored repos. Automated, takes minutes.',
    'Run doc generation per section. LLM produces output in about 30 minutes per section.',
    'First draft is available within 1 day of feature freeze.',
    'Docs team reviews and edits over the following months.',
    'Re-generate as needed if late code changes land.',
    'GA release: docs ship alongside the product.',
    'Post-GA: evaluate generated vs final published docs, update SKILL.md for the next cycle.',
]
for i, s in enumerate(workflow_steps, 1):
    p = doc.add_paragraph()
    p.add_run(f'{i}. ').bold = True
    p.add_run(s)

doc.add_heading('One PR per section', level=2)
doc.add_paragraph(
    'Instead of one massive PR covering all doc changes, each section gets its own branch and PR. '
    'This makes it easier for the docs team to review and merge incrementally.'
)

add_table(
    ['Section', 'Branch', 'PR Target'],
    [
        ['installing', 'ai-docs/installing-4.23', 'enterprise-4.23'],
        ['updating', 'ai-docs/updating-4.23', 'enterprise-4.23'],
        ['networking (future)', 'ai-docs/networking-4.23', 'enterprise-4.23'],
    ]
)

doc.add_paragraph()
doc.add_heading('Scaling to more sections', level=2)
doc.add_paragraph('The same approach applies to any documentation section. To add a new section:')

scale_steps = [
    'Identify the relevant source repositories for that section.',
    'Create a SKILL.md with section-specific rules (start with about 12 generic rules).',
    'Run the training loop against versions 4.16 through 4.22 to grow the rule set.',
    'The section is now ready for production generation.',
]
for s in scale_steps:
    doc.add_paragraph(s, style='List Number')

doc.add_paragraph()
doc.add_paragraph('Candidate sections for expansion: networking, security, observability, CI/CD, virtualization.')

# ============================================================
# 12. PROJECT SUMMARY
# ============================================================
doc.add_heading('12. Project Summary', level=1)

add_table(
    ['Metric', 'Value'],
    [
        ['Sections trained', '2 (Installing, Updating)'],
        ['Total iterations run', '12 (6 per section)'],
        ['Total rules learned', '51 (23 + 28)'],
        ['Source repos (installing)', '7 (installer, api, baremetal-operator, assisted-installer, cluster-network-operator, machine-config-operator, machine-api-operator)'],
        ['Source repos (updating)', '5 (cluster-version-operator, oc, machine-config-operator, api, cluster-network-operator)'],
        ['Unique repos across both', '9'],
        ['Avg deterministic accuracy', '94.8% (installing), 96.5% (updating)'],
        ['Best accuracy achieved', '98.8% (updating 4.21)'],
        ['LLM completeness (enhanced)', '80.8% (installing 4.17)'],
        ['Time to first draft', 'Less than 1 day after feature freeze'],
        ['Hard ceiling (code-only)', 'About 75 to 80% of all doc changes'],
        ['With JIRA context (projected)', 'About 90% of all doc changes'],
    ]
)

doc.add_paragraph()
doc.add_heading('Links', level=2)

links = [
    ('GitHub Repository', 'https://github.com/sarthakpurohit/ai-user-docs'),
    ('Interactive Presentation', 'presentation/project-overview-final.html (in repo, open in browser)'),
    ('Installing SKILL.md', 'skills/generate-install-docs/SKILL.md'),
    ('Updating SKILL.md', 'skills/generate-updating-docs/SKILL.md'),
]
for name, url in links:
    p = doc.add_paragraph()
    p.add_run(f'{name}: ').bold = True
    p.add_run(url)

doc.add_paragraph()
doc.add_heading('How to try the comparison viewer', level=2)
doc.add_paragraph('See the AI-generated docs compared against human-written ones:')

try_steps = [
    'Clone the repo: git clone https://github.com/sarthakpurohit/ai-user-docs.git',
    'Extract ground truth: make extract SECTION=installing',
    'Launch viewer: make compare VERSION=4.17 SECTION=installing',
    'A browser opens with 3-panel side-by-side comparison (Previous | Human-written | AI-generated).',
]
for s in try_steps:
    doc.add_paragraph(s, style='List Number')

# ============================================================
# SAVE
# ============================================================
outpath = os.path.expanduser(
    '~/Desktop/AI-Generated User-Facing Documentation - Project Overview.docx'
)
doc.save(outpath)
print(f'Saved: {outpath}')
